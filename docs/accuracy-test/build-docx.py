# -*- coding: utf-8 -*-
"""
Perakit naskah jurnal Angle A -> Word (.docx) siap cetak.
Menggabungkan: pendahuluan-tinjauan-pustaka.md + paper-draft.md + simpulan-abstrak.md
(tanpa catatan penyuntingan), tabel dinomori ulang 1-9, gambar 1-6 tertanam.

  python docs/accuracy-test/build-docx.py
  -> docs/accuracy-test/naskah-jurnal-angle-a.docx
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(HERE, 'results', 'figures')
OUT = os.path.join(HERE, 'naskah-jurnal-angle-a.docx')

doc = Document()

# ---------- halaman & gaya dasar ----------
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)          # A4
sec.top_margin = sec.bottom_margin = Cm(2.5)
sec.left_margin, sec.right_margin = Cm(2.5), Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15


def para(text='', size=11, bold=False, italic=False, align='justify',
         indent=True, color=None, space_after=6):
    p = doc.add_paragraph()
    p.alignment = {
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'left': WD_ALIGN_PARAGRAPH.LEFT,
    }[align]
    if indent and align == 'justify':
        p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p


def heading(num, text, level=1):
    label = f'{num}. {text}' if num else text
    p = para(label, size=12 if level == 1 else 11, bold=True,
             align='left', indent=False, space_after=6)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    return p


def caption(text, size=10):
    return para(text, size=size, align='center', indent=False, space_after=4)


def table(headers, rows, widths=None, font=10):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ''
        r = cell.paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.size = Pt(font)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = ''
            r = cell.paragraphs[0].add_run(str(v))
            r.font.size = Pt(font)
            cell.paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER)
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Cm(w)
    para('', size=4, space_after=4, indent=False)  # jarak setelah tabel
    return t


def figure(png, cap):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(os.path.join(FIG, png), width=Cm(15))
    caption(cap)


def equation(text, num):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.italic = True
    r.font.size = Pt(11)
    r2 = p.add_run(f'    ({num})')
    r2.font.size = Pt(11)


# ================= BLOK PERINGATAN DRAF =================
para('CATATAN DRAF — HAPUS SEBELUM SUBMIT: (1) isi identitas penulis; '
     '(2) verifikasi seluruh referensi [1]–[15] ke sumber asli (DOI, penulis, halaman); '
     '(3) sunting dengan gaya bahasa Anda sendiri; (4) sesuaikan dengan template jurnal target; '
     '(5) periksa kebijakan penggunaan AI jurnal target.',
     size=10, bold=True, color=(0xC0, 0x00, 0x00), align='left', indent=False, space_after=12)

# ================= JUDUL & PENULIS =================
para('Pengujian Akurasi Algoritma Geospasial untuk Deteksi Ritual Haji '
     'Berbasis GPS terhadap Galat Posisi',
     size=14, bold=True, align='center', indent=False, space_after=10)
para('[Nama Penulis Pertama]1*, [Nama Penulis Kedua]2',
     size=11, align='center', indent=False, space_after=2)
para('1,2[Program Studi, Institusi — mis. STMIK Tazkia, Bogor]',
     size=10, align='center', indent=False, space_after=2)
para('*e-mail korespondensi: [nama@stmik.tazkia.ac.id]',
     size=10, italic=True, align='center', indent=False, space_after=14)

# ================= ABSTRAK =================
para('ABSTRAK', size=11, bold=True, align='center', indent=False)
para('Keabsahan ibadah haji terikat pada lokasi dan hitungan: wukuf harus di dalam batas Arafah, '
     'tawaf genap tujuh putaran, dan sa’i tujuh lintasan dari Safa ke Marwah. Aplikasi pendamping haji '
     'yang ada umumnya berhenti pada penyajian konten manasik, sementara sistem yang memanfaatkan lokasi '
     'belum pernah diuji ketepatan algoritmanya terhadap galat GPS ponsel yang mencapai 7–13 m di lingkungan '
     'padat bangunan. Penelitian ini merancang enam algoritma deteksi ritual berbasis geometri posisi — '
     'perhitungan jarak haversine, geofence miqat, deteksi batas Arafah dengan ray casting, penghitung tawaf '
     'berbasis persilangan sudut, penghitung sa’i berbasis pergantian zona, dan identifikasi pilar jamarat — '
     'lalu menguji akurasinya melalui simulasi Monte Carlo dengan injeksi derau Gaussian bertingkat '
     '(σ = 1–15 m) yang sepenuhnya dapat direproduksi. Hasilnya: penghitung sa’i bertahan 100% pada seluruh '
     'tingkat galat; geofence miqat dan deteksi Arafah di atas 99,4%; penghitung tawaf andal hingga σ = 10 m '
     '(100%) sebelum turun ke 72,67% pada σ = 15 m; identifikasi jamarat akurat hingga σ = 5 m dengan '
     'kesalahan antar-pilar maksimum 0,32%. Ketahanan algoritma terbukti ditentukan oleh rasio separasi '
     'geometris objek ritual terhadap besar galat posisi, temuan yang dapat menjadi dasar perancangan '
     'aplikasi ibadah berbasis lokasi. Pada rentang galat GPS ponsel yang tipikal, keenam algoritma berada '
     'dalam wilayah kerja yang aman tanpa memerlukan model pembelajaran mesin maupun sensor tambahan.',
     size=10, indent=False)
para('Kata kunci: deteksi ritual haji; geofencing; akurasi GPS; simulasi Monte Carlo; haversine; point-in-polygon',
     size=10, italic=True, indent=False, space_after=12)

para('ABSTRACT', size=11, bold=True, align='center', indent=False)
para('The validity of Hajj rituals is bound to location and count: wukuf must take place within the '
     'boundaries of Arafah, tawaf requires exactly seven circuits, and sa’i seven laps from Safa to Marwah. '
     'Existing Hajj companion applications mostly stop at presenting ritual guidance content, while '
     'location-aware systems have never had their algorithms tested against smartphone GPS errors, which '
     'reach 7–13 m in built-up environments. This study designs six position-geometry-based ritual detection '
     'algorithms — haversine distance computation, miqat geofencing, Arafah boundary detection using ray '
     'casting, an angular-crossing tawaf counter, a zone-alternation sa’i counter, and jamarat pillar '
     'identification — and evaluates their accuracy through fully reproducible Monte Carlo simulation with '
     'graduated Gaussian noise injection (σ = 1–15 m). The results show three tiers of robustness: the sa’i '
     'counter remains perfect (100%) at all noise levels; miqat geofencing and Arafah detection stay above '
     '99.4%; the tawaf counter is fully reliable up to σ = 10 m (100%) before dropping to 72.67% at '
     'σ = 15 m; and jamarat identification is accurate up to σ = 5 m with cross-pillar misidentification '
     'never exceeding 0.32%. Algorithm robustness is shown to be governed by the ratio of the ritual '
     'objects’ geometric separation to the magnitude of positioning error — a finding that can inform the '
     'design of similar location-based worship applications. Within the typical range of smartphone GPS '
     'error, all six algorithms operate safely without requiring machine learning models or additional sensors.',
     size=10, italic=True, indent=False)
para('Keywords: Hajj ritual detection; geofencing; GPS accuracy; Monte Carlo simulation; haversine; point-in-polygon',
     size=10, italic=True, indent=False, space_after=12)

# ================= 1. PENDAHULUAN =================
heading('1', 'PENDAHULUAN')
para('Indonesia memberangkatkan jamaah haji dalam jumlah yang tidak dimiliki negara mana pun. '
     'Musim haji 1445 H/2024 M mencatat 241.000 jamaah, kuota terbesar sepanjang sejarah penyelenggaraan '
     'haji Indonesia, dan sekitar 45.000 di antaranya berusia 65 tahun ke atas [1], [2]. Di balik angka itu '
     'ada persoalan yang berulang hampir setiap musim: jamaah tersesat. Laporan penyelenggaraan haji 2024 '
     'mencatat kasus jamaah lanjut usia yang kehilangan arah karena minimnya penunjuk jalan, dan Kementerian '
     'Agama sendiri menempatkan mitigasi risiko jamaah lansia sebagai catatan perbaikan layanan [2]. '
     'Persoalan ini tidak sederhana. Kawasan Masjidil Haram, Arafah, Muzdalifah, dan Mina pada puncak musim '
     'haji dipadati jutaan manusia, sementara seorang pembimbing (muthawwif) lazimnya harus mengawal '
     'puluhan jamaah sekaligus.')
para('Ada satu karakter ibadah haji yang jarang disentuh aplikasi pendamping pada umumnya: keabsahannya '
     'terikat pada lokasi dan hitungan. Wukuf hanya sah bila dilakukan di dalam batas wilayah Arafah — dan '
     'batas ini bukan perkara sepele, karena sebagian bangunan Masjid Namirah justru berada di luarnya. '
     'Tawaf harus genap tujuh putaran mengelilingi Ka’bah, dihitung dari garis sejajar Hajar Aswad. Sa’i '
     'harus tujuh lintasan yang dimulai dari Safa dan berakhir di Marwah. Ihram harus diniatkan sebelum '
     'melintasi miqat. Bagi jamaah lansia yang kelelahan di tengah kepadatan, sekadar mengingat “ini putaran '
     'ke berapa” adalah beban kognitif tersendiri; salah hitung berarti keraguan terhadap sahnya ibadah yang '
     'telah dinanti puluhan tahun.')
para('Aplikasi pendamping haji dan umrah yang tersedia saat ini, maupun yang dilaporkan dalam literatur, '
     'umumnya berhenti pada penyajian konten: materi manasik, kumpulan doa, dan jadwal kegiatan [3], [4], [5]. '
     'Beberapa penelitian telah melangkah ke pemantauan posisi jamaah berbasis geofencing [6], namun '
     'evaluasinya terbatas pada penerimaan pengguna melalui kuesioner, bukan pada ketepatan algoritmanya. '
     'Di sisi lain, penelitian deteksi ritual secara otomatis mulai bermunculan dengan pendekatan sensor '
     'inersia (IMU) [7], [8] dan visi komputer [9]. Pendekatan-pendekatan tersebut menjanjikan, tetapi '
     'membawa prasyarat yang berat untuk konteks jamaah Indonesia: model pembelajaran mesin yang harus '
     'dilatih, konsumsi komputasi yang tinggi, atau infrastruktur kamera yang hanya dimiliki otoritas '
     'setempat. Padahal setiap jamaah sudah menggenggam satu sensor posisi: GPS pada ponselnya.')
para('Pertanyaannya, dapatkah deteksi ritual dibangun hanya dari geometri posisi GPS — tanpa pelatihan '
     'model, tanpa sensor tambahan — dan seberapa jauh ia bertahan terhadap galat pengukuran? Pertanyaan '
     'kedua ini yang justru paling menentukan dan paling jarang dijawab. Studi empiris menunjukkan galat '
     'horizontal GPS ponsel berkisar 5–8,5 m pada kondisi baik [10] dan mencapai 7–13 m di lingkungan '
     'perkotaan yang padat bangunan [11]. Lingkungan sekitar Masjidil Haram, yang dikelilingi menara-menara '
     'tinggi, termasuk kategori terburuk untuk penerimaan sinyal satelit. Sebuah algoritma penghitung tawaf '
     'yang bekerja sempurna pada koordinat ideal bisa saja gagal total pada koordinat yang bergeser '
     'sepuluh meter.')
para('Penelitian ini merancang dan menguji enam algoritma geospasial untuk deteksi ritual haji dan umrah '
     'pada sistem Mabrur, sebuah aplikasi pendamping jamaah berbasis Android: (i) perhitungan jarak '
     'haversine, (ii) geofence miqat, (iii) deteksi batas Arafah dengan poligon ray casting, (iv) penghitung '
     'putaran tawaf berbasis persilangan sudut, (v) penghitung lintasan sa’i berbasis pergantian zona, dan '
     '(vi) identifikasi pilar jamarat. Kontribusi utamanya bukan pada masing-masing formula — haversine dan '
     'ray casting adalah metode mapan — melainkan pada tiga hal: pertama, komposisi keenamnya menjadi satu '
     'rangkaian deteksi ritual yang berjalan sepenuhnya di perangkat tanpa model terlatih; kedua, kerangka '
     'pengujian akurasi berbasis simulasi Monte Carlo dengan injeksi derau GPS Gaussian bertingkat '
     '(σ = 1–15 m) yang seluruhnya dapat direproduksi; ketiga, temuan kuantitatif tentang hubungan antara '
     'separasi geometris objek ritual dan ketahanan algoritma terhadap galat posisi, yang dapat menjadi '
     'dasar perancangan aplikasi sejenis.')
para('Sisa naskah ini disusun sebagai berikut. Bagian 2 mengulas penelitian terkait dan landasan metode '
     'geometris yang digunakan. Bagian 3 memaparkan rancangan algoritma serta prosedur pengujian akurasinya. '
     'Bagian 4 menyajikan hasil pengujian beserta pembahasannya, dan Bagian 5 menutup dengan simpulan serta '
     'arah penelitian lanjutan.')

# ================= 2. TINJAUAN PUSTAKA =================
heading('2', 'TINJAUAN PUSTAKA')
heading('2.1', 'Aplikasi Pendamping Haji dan Umrah', level=2)
para('Penelitian aplikasi haji dan umrah di Indonesia selama satu dekade terakhir dapat dikelompokkan '
     'menurut fungsi yang ditawarkannya. Kelompok pertama, dan yang paling banyak, adalah aplikasi '
     'pembelajaran manasik: penyajian tata cara ibadah, bacaan doa, dan materi bimbingan dalam bentuk '
     'digital sebagai pengganti buku saku [3], [4]. Kelompok kedua adalah aplikasi panduan perjalanan, '
     'seperti pemodelan Umrah Guide yang menstrukturkan informasi syarat, larangan ihram, dan prosedur '
     'umrah [5]. Kedua kelompok ini pada dasarnya memindahkan konten cetak ke layar ponsel; posisi jamaah '
     'tidak dilibatkan dalam logika aplikasi.')
para('Kelompok ketiga mulai memanfaatkan lokasi. Budiawan dkk. mengembangkan aplikasi pemantauan jamaah '
     'berbasis geofencing dan Firebase Cloud Messaging untuk sebuah biro perjalanan, yang memungkinkan '
     'pembimbing memantau sebaran jamaahnya dan jamaah meminta bantuan [6]. Evaluasi yang dilaporkan berupa '
     'tingkat penerimaan: 93,33% pembimbing menilai aplikasi membantu pemantauan dan 86% jamaah menilainya '
     'bermanfaat. Angka tersebut menunjukkan kebutuhan yang nyata, namun menyisakan pertanyaan yang belum '
     'dijawab: seberapa tepat geofence itu sendiri bekerja, dan berapa banyak keputusan masuk/keluar zona '
     'yang salah akibat galat posisi. Sepanjang penelusuran penulis, belum ada penelitian aplikasi haji di '
     'Indonesia yang mengevaluasi akurasi algoritma lokasinya secara kuantitatif.')
heading('2.2', 'Deteksi Ritual Secara Otomatis', level=2)
para('Di tingkat internasional, deteksi aktivitas ritual haji telah dicoba melalui beberapa jalur. Jalur '
     'pertama memakai sensor inersia. Kammoun dkk. mendeteksi aktivitas tawaf dan sa’i dari data '
     'akselerometer dan giroskop ponsel yang diproses dengan algoritma pengenalan aktivitas [7]; pendekatan '
     'serupa dikembangkan menjadi model penghitung putaran tawaf dan sa’i yang berjalan mandiri di '
     'perangkat [8]. Jalur kedua memakai visi komputer: jaringan saraf konvolusional untuk mengestimasi '
     'kepadatan dan menghitung individu pada rekaman video area tawaf, dengan akurasi yang dilaporkan '
     'sekitar 87% pada dataset HAJJv2 [9].')
para('Kedua jalur tersebut memiliki wilayah keunggulannya masing-masing, tetapi juga batasan yang jelas. '
     'Pendekatan IMU peka terhadap variasi cara membawa ponsel dan gaya berjalan — faktor yang sulit '
     'dikendalikan pada jamaah lansia yang memakai kursi roda atau skuter, yang jumlahnya justru terus '
     'bertambah. Pendekatan visi membutuhkan infrastruktur kamera dan komputasi yang berada di luar kendali '
     'penyelenggara ibadah dari negara pengirim. Pendekatan geometris berbasis GPS, yang dipakai penelitian '
     'ini, menukar kecanggihan model dengan kesederhanaan: tidak ada pelatihan, tidak ada sensor tambahan, '
     'logikanya dapat diaudit baris per baris — kualitas yang relevan untuk perangkat lunak yang menyentuh '
     'keabsahan ibadah. Konsekuensinya, seluruh beban ketepatan berpindah ke satu titik: kualitas posisi '
     'GPS. Karena itu pengujian ketahanan terhadap galat posisi bukan pelengkap, melainkan inti dari '
     'validasi pendekatan ini.')
heading('2.3', 'Karakteristik Galat Posisi GPS pada Ponsel', level=2)
para('Besaran galat GPS ponsel telah diukur pada berbagai kondisi. Zandbergen dan Barbeau mencatat median '
     'galat horizontal 5,0–8,5 m pada pengujian statis di ruang terbuka, dua sampai tiga kali lebih besar '
     'dibanding penerima GPS khusus [10]. Merry dan Bettinger mengukur galat 7–13 m di lingkungan perkotaan '
     '[11], dan studi-studi yang lebih baru pada berbagai merek ponsel melaporkan rentang 2–15 m terhadap '
     'posisi referensi diferensial [12]. Kawasan Masjidil Haram menghadirkan kombinasi terburuk dari '
     'kondisi-kondisi tersebut: gedung menjulang di sekelilingnya memantulkan sinyal satelit (efek '
     'multipath), sementara kepadatan manusia menghalangi pandangan langit. Rentang inilah yang menjadi '
     'dasar pemilihan tingkat derau σ = 1 sampai 15 m pada pengujian di Bagian 3: nilai kecil mewakili '
     'kondisi ideal ruang terbuka, nilai besar mewakili skenario pesimistis di sekitar bangunan tinggi.')
heading('2.4', 'Metode Geometris yang Menjadi Landasan', level=2)
para('Tiga metode klasik menjadi fondasi algoritma yang diuji. Formula haversine menghitung jarak lingkaran '
     'besar antara dua koordinat dengan asumsi bumi berbentuk bola [13]; formula ini stabil secara numerik '
     'untuk jarak pendek — wilayah kerja utama deteksi ritual — namun membawa galat sistematik terhadap '
     'bentuk bumi yang sesungguhnya elipsoid. Formula Vincenty menghitung jarak pada elipsoid WGS-84 secara '
     'iteratif dengan ketelitian milimeter [14], sehingga lazim dipakai sebagai acuan pembanding; peran itu '
     'pula yang diberikan kepadanya dalam penelitian ini. Untuk menentukan apakah sebuah titik berada di '
     'dalam poligon, algoritma ray casting menarik garis dari titik uji dan menghitung paritas '
     'perpotongannya dengan sisi-sisi poligon [15]; kesederhanaannya menjadikan metode ini pilihan umum '
     'untuk geofence berbentuk bebas seperti batas wilayah Arafah.')
heading('2.5', 'Posisi Penelitian Ini', level=2)
para('Rangkuman di atas memperlihatkan celah yang hendak diisi. Penelitian aplikasi haji di dalam negeri '
     'kuat pada sisi konten dan mulai menyentuh pemanfaatan lokasi, tetapi berhenti pada evaluasi penerimaan '
     'pengguna [3]–[6]. Penelitian deteksi ritual di luar negeri kuat pada sisi algoritma, tetapi bertumpu '
     'pada IMU atau visi komputer dengan prasyarat yang berat [7]–[9], dan belum ada yang menguji ketahanan '
     'pendekatannya terhadap galat posisi secara sistematis. Penelitian ini berdiri di persilangan keduanya: '
     'algoritma deteksi ritual berbasis geometri GPS murni yang dievaluasi bukan dengan kuesioner, melainkan '
     'dengan pengujian akurasi terkuantifikasi pada enam tingkat galat — disertai kode dan prosedur yang '
     'dapat direproduksi penuh.')

# ================= 3. METODE PENGUJIAN =================
heading('3', 'METODE PENGUJIAN')
heading('3.1', 'Rancangan Eksperimen', level=2)
para('Pengujian akurasi enam algoritma geospasial inti dilakukan melalui simulasi Monte Carlo. Pendekatan '
     'simulasi dipilih karena pengujian lapangan langsung di area Masjidil Haram, Arafah, dan Jamarat tidak '
     'dapat dilaksanakan secara praktis. Prinsip pengujian: setiap algoritma dijalankan dengan kode sumber '
     'asli aplikasi, lalu diberi masukan berupa posisi ground truth yang telah diketahui secara geometris, '
     'dan diamati apakah keluaran algoritma tetap benar ketika posisi masukan diberi gangguan (derau) GPS. '
     'Keenam algoritma yang diuji dirangkum pada Tabel 1.')
caption('Tabel 1. Enam algoritma yang diuji')
table(
    ['No', 'Algoritma', 'Teknik', 'Peran dalam sistem'],
    [
        ['1', 'Haversine', 'Great-circle distance', 'Dasar seluruh perhitungan jarak'],
        ['2', 'Geofence Miqat', 'Point-in-circle', 'Peringatan batas miqat & ihram'],
        ['3', 'Deteksi Arafah', 'Ray-casting polygon', 'Validasi keabsahan wukuf'],
        ['4', 'Penghitung Tawaf', 'Angular crossing detection', 'Hitung otomatis 7 putaran'],
        ['5', "Penghitung Sa'i", 'Zone alternation', "Hitung otomatis 7 leg Safa–Marwah"],
        ['6', 'Deteksi Jamarat', 'Nearest-in-radius (3 kelas)', 'Identifikasi pilar lempar jumrah'],
    ], widths=[1.0, 3.6, 4.8, 6.2])
heading('3.2', 'Model Gangguan (Derau) GPS', level=2)
para('Ketidakpastian posisi GPS pada perangkat ponsel dimodelkan sebagai gangguan Gaussian isotropik yang '
     'ditambahkan pada komponen Timur (East) dan Utara (North) posisi sebenarnya:')
equation('posisi_terukur = posisi_sebenarnya + N(0, σ) pada tiap sumbu (E, N)', 1)
para('dengan σ adalah simpangan baku galat posisi (meter). Konversi meter ke derajat menggunakan '
     '1° lintang ≈ 111.320 m dan 1° bujur ≈ 111.320 × cos(lintang) m. Pengujian menggunakan lima tingkat '
     'galat yang mengacu pada rentang akurasi GPS ponsel di ruang terbuka hingga lingkungan padat bangunan '
     '[10]–[12]: σ ∈ {1, 3, 5, 10, 15} meter, ditambah σ = 0 sebagai baseline untuk memverifikasi '
     'kebenaran algoritma tanpa gangguan.')
heading('3.3', 'Variabel dan Metrik', level=2)
para('Variabel bebas pengujian adalah tingkat galat GPS (σ); variabel terikatnya adalah metrik akurasi '
     'tiap algoritma. Untuk algoritma klasifikasi (Geofence Miqat, Deteksi Arafah, Deteksi Jamarat) '
     'digunakan confusion matrix dengan metrik:')
equation('Akurasi = (TP + TN) / (TP + TN + FP + FN)', 2)
equation('Presisi = TP / (TP + FP)', 3)
equation('Recall = TP / (TP + FN)', 4)
equation('F1 = 2 × (Presisi × Recall) / (Presisi + Recall)', 5)
para('Untuk algoritma pengukuran dan penghitungan (Haversine, Tawaf, Sa’i) digunakan galat terhadap '
     'nilai acuan:')
equation('MAE = (1/n) Σ |ŷ − y|', 6)
equation('RMSE = √[(1/n) Σ (ŷ − y)²]', 7)
para('Akurasi Haversine diukur terhadap formula Vincenty (model elipsoid WGS-84) sebagai acuan presisi '
     'tinggi. Akurasi Tawaf dan Sa’i diukur sebagai counting accuracy, yaitu persentase percobaan yang '
     'menghasilkan hitungan tepat sama dengan tujuh.')
heading('3.4', 'Prosedur dan Reproduksibilitas', level=2)
para('Bilangan acak dibangkitkan dengan generator mulberry32 ber-seed tetap (seed = 42), sehingga seluruh '
     'hasil dapat direproduksi secara identik. Jumlah sampel: 5.000 pasang titik (Haversine), 8.000 titik '
     'per σ (Miqat), 12.000 titik per σ (Arafah), 300 percobaan lintasan per σ (Tawaf dan Sa’i), dan '
     '12.000 titik per σ (Jamarat). Lintasan Tawaf disimulasikan sebagai gerak melingkar berjari-jari 25 m '
     '(sekitar 0,52 m/s, laju tawaf padat), lintasan Sa’i sebagai gerak bolak-balik Safa–Marwah, keduanya '
     'dengan pencuplikan 3 detik sesuai mode BestForNavigation aplikasi. Parameter geometris yang dihitung '
     'dari koordinat sistem dirangkum pada Tabel 2.')
caption('Tabel 2. Parameter geometris objek ritual')
table(
    ['Besaran', 'Nilai'],
    [
        ['Jarak Safa–Marwah', '419,0 m'],
        ['Jarak antar-Jamarat (Ula–Wustha / Wustha–Aqabah / Ula–Aqabah)', '76,0 / 68,2 / 144,0 m'],
        ['Radius deteksi Jamarat', '30 m'],
        ["Radius zona Sa'i (Safa/Marwah)", '25 m'],
        ["Band radius Tawaf dari Ka'bah", '10–80 m'],
    ], widths=[9.5, 6.0])

# ================= 4. HASIL DAN PEMBAHASAN =================
heading('4', 'HASIL DAN PEMBAHASAN')
heading('4.1', 'Akurasi Haversine terhadap Elipsoid WGS-84', level=2)
caption('Tabel 3. Galat Haversine vs Vincenty')
table(
    ['Skenario jarak', 'MAE (m)', 'RMSE (m)', 'Error rata-rata (%)', 'Error maks (%)'],
    [
        ['Lokal Masjidil Haram (0–0,5 km)', '0,649', '0,811', '0,2000', '0,4267'],
        ['Skala Miqat (10–450 km)', '521,533', '675,987', '0,2654', '0,4270'],
    ], widths=[6.0, 2.3, 2.3, 2.7, 2.3])
para('Formula Haversine mengasumsikan bumi berbentuk bola sempurna (R = 6.371 km), sedangkan Vincenty '
     'memodelkan elipsoid. Galat relatif tetap di bawah 0,5% pada kedua skala. Untuk deteksi ritual '
     'berskala meter (tawaf, sa’i, jamarat), galat absolut kurang dari 1 m tidak berdampak signifikan. '
     'Pada skala miqat, galat absolut ratusan meter tampak besar namun masih sekitar 0,27% dan berada jauh '
     'di bawah radius peringatan miqat 3.000 m, sehingga tidak mengubah keputusan geofence.')
heading('4.2', 'Geofence Miqat (Point-in-Circle)', level=2)
caption('Tabel 4. Akurasi klasifikasi “dalam batas miqat” (radius 1.000 m) terhadap σ')
table(
    ['σ (m)', 'Akurasi (%)', 'Presisi (%)', 'Recall (%)', 'F1 (%)'],
    [
        ['0', '100,00', '100,00', '100,00', '100,00'],
        ['1', '100,00', '100,00', '100,00', '100,00'],
        ['3', '99,92', '99,87', '99,94', '99,91'],
        ['5', '99,81', '99,68', '99,84', '99,76'],
        ['10', '99,69', '99,59', '99,62', '99,61'],
        ['15', '99,45', '99,49', '99,11', '99,30'],
    ])
para('Akurasi menurun sangat landai, tetap di atas 99,4% bahkan pada σ = 15 m (Gambar 1). Kesalahan hanya '
     'terjadi pada titik yang jaraknya sangat dekat dengan garis batas 1.000 m, di mana gangguan GPS dapat '
     'membalik klasifikasi. Karena radius peringatan aktual sistem (3.000 m) jauh lebih longgar, keandalan '
     'praktis fitur peringatan miqat sangat tinggi.')
figure('fig2-miqat-metrik.png',
       'Gambar 1. Metrik klasifikasi Geofence Miqat (akurasi, presisi, recall, F1) terhadap σ; '
       'seluruh metrik tetap di atas 99,1%.')
heading('4.3', 'Deteksi Arafah (Ray-Casting Polygon)', level=2)
caption('Tabel 5. Akurasi klasifikasi dalam/luar poligon Arafah terhadap σ')
table(
    ['σ (m)', 'Akurasi (%)', 'Presisi (%)', 'Recall (%)', 'F1 (%)'],
    [
        ['0', '100,00', '100,00', '100,00', '100,00'],
        ['1', '99,98', '100,00', '99,96', '99,98'],
        ['3', '99,95', '99,99', '99,93', '99,96'],
        ['5', '99,84', '99,82', '99,90', '99,86'],
        ['10', '99,63', '99,68', '99,68', '99,68'],
        ['15', '99,44', '99,48', '99,55', '99,52'],
    ])
para('Algoritma ray-casting mengklasifikasikan posisi dengan benar di atas 99,4% pada seluruh tingkat '
     'gangguan (Gambar 2). Kesalahan terkonsentrasi pada pita sempit di sekitar tepi poligon (selebar '
     'kira-kira σ); titik yang jauh di dalam atau di luar Arafah selalu diklasifikasikan benar. Ini relevan '
     'untuk validasi keabsahan wukuf, di mana kesalahan hanya mungkin bagi jamaah yang benar-benar berdiri '
     'tepat di garis batas.')
figure('fig3-arafah-metrik.png',
       'Gambar 2. Metrik klasifikasi Deteksi Arafah (ray-casting polygon) terhadap σ; kesalahan '
       'terkonsentrasi pada pita tepi poligon.')
heading('4.4', 'Penghitung Tawaf Otomatis', level=2)
caption('Tabel 6. Counting accuracy Tawaf (target = 7 putaran, 300 percobaan per σ)')
table(
    ['σ (m)', 'Rata-rata putaran', 'Akurasi tepat-7 (%)', 'MAE', 'RMSE'],
    [
        ['0', '7,00', '100,00', '0,000', '0,000'],
        ['1', '7,00', '100,00', '0,000', '0,000'],
        ['3', '7,00', '100,00', '0,000', '0,000'],
        ['5', '7,00', '100,00', '0,000', '0,000'],
        ['10', '7,00', '100,00', '0,000', '0,000'],
        ['15', '7,31', '72,67', '0,313', '0,632'],
    ])
para('Penghitung tawaf sangat andal hingga σ = 10 m (akurasi 100%), lalu menurun ke 72,67% pada σ = 15 m '
     '(Gambar 3). Pada gangguan besar, posisi ber-derau di sekitar garis Hajar Aswad dapat memicu deteksi '
     'persilangan ganda; rata-rata putaran di atas tujuh (7,31) menunjukkan kecenderungan kelebihan hitung '
     '(over-count). Mekanisme debounce 120 detik meredam sebagian galat ini namun tidak sepenuhnya '
     'menghilangkannya pada radius tawaf yang kecil (25 m).')
figure('fig4-tawaf.png',
       'Gambar 3. Akurasi tepat-7 penghitung Tawaf otomatis; andal hingga σ = 10 m (100%), menurun ke '
       '72,67% pada σ = 15 m akibat over-count di sekitar garis Hajar Aswad.')
heading('4.5', "Penghitung Sa'i Otomatis", level=2)
caption("Tabel 7. Counting accuracy Sa'i (target = 7 leg, 300 percobaan per σ)")
table(
    ['σ (m)', 'Rata-rata leg', 'Akurasi tepat-7 (%)', 'MAE', 'RMSE'],
    [['0–15', '7,00', '100,00', '0,000', '0,000']])
para('Penghitung sa’i sempurna (100%) pada seluruh tingkat gangguan, termasuk σ = 15 m (Gambar 4). '
     'Ketahanan ini disebabkan pemisahan geometris Safa–Marwah yang besar (419 m) — jauh melampaui skala '
     'gangguan GPS — sehingga transisi antar-zona selalu terdeteksi tanpa ambiguitas. Ini menunjukkan bahwa '
     'keandalan penghitung berbasis zona berbanding lurus dengan rasio jarak antar-zona terhadap galat GPS.')
figure('fig5-sai.png',
       "Gambar 4. Akurasi tepat-7 penghitung Sa'i otomatis; 100% pada seluruh σ berkat separasi geometris "
       'Safa–Marwah (419 m) yang jauh melampaui galat GPS. Skala sumbu-y disamakan dengan Gambar 3 untuk '
       'perbandingan.')
heading('4.6', 'Deteksi Jamarat (Klasifikasi 3 Kelas)', level=2)
caption('Tabel 8. Akurasi identifikasi pilar Jamarat terhadap σ')
table(
    ['σ (m)', 'Akurasi benar (%)', 'Salah pilar (%)', 'Tak terdeteksi (%)'],
    [
        ['0', '100,00', '0,00', '0,00'],
        ['1', '100,00', '0,00', '0,00'],
        ['3', '100,00', '0,00', '0,00'],
        ['5', '100,00', '0,00', '0,00'],
        ['10', '97,31', '0,03', '2,67'],
        ['15', '83,53', '0,32', '16,15'],
    ])
caption('Tabel 9. Confusion matrix Jamarat pada σ = 15 m')
table(
    ['Sebenarnya \\ Prediksi', 'Ula', 'Wustha', 'Aqabah', 'Tak terdeteksi'],
    [
        ['Ula', '3.328', '4', '0', '668'],
        ['Wustha', '2', '3.364', '18', '616'],
        ['Aqabah', '0', '14', '3.332', '654'],
    ])
para('Deteksi jamarat akurat 100% hingga σ = 5 m, lalu menurun ke 83,53% pada σ = 15 m (Gambar 5). Yang '
     'penting: kesalahan salah-pilar hampir nol (maksimum 0,32%) — karena jarak antar-pilar (68–144 m) '
     'melebihi dua kali radius deteksi (60 m), sehingga zona tidak tumpang tindih. Degradasi hampir '
     'seluruhnya berupa “tak terdeteksi” (16,15% pada σ = 15 m): gangguan besar mendorong posisi terukur '
     'keluar dari radius 30 m. Implikasinya, risiko utama pada gangguan tinggi bukan salah identifikasi '
     'melainkan gagal deteksi, yang dapat dimitigasi dengan memperbesar radius atau menggabungkan konteks '
     'urutan lempar per hari.')
figure('fig6-jamarat-hasil.png',
       'Gambar 5. Komposisi hasil deteksi Jamarat terhadap σ: kesalahan salah-pilar nyaris nol (≤ 0,32%); '
       'degradasi pada σ besar didominasi kegagalan deteksi (posisi terukur keluar radius 30 m).')
heading('4.7', 'Ringkasan Ketahanan Antar-Algoritma', level=2)
para('Tingkat ketahanan terhadap gangguan GPS berkorelasi dengan rasio separasi geometris terhadap σ '
     '(Gambar 6). Sa’i, dengan separasi 419 m, paling tahan — 100% pada semua σ. Geofence Miqat dan '
     'Arafah, yang keputusannya hanya sensitif di pita tepi, bertahan di atas 99,4% pada semua σ. Tawaf, '
     'dengan radius hanya 25 m, andal hingga σ = 10 m lalu menurun. Jamarat, dengan radius deteksi 30 m, '
     'andal hingga σ = 5 m dan menurun melalui non-deteksi pada σ ≥ 10 m.')
figure('fig1-gabungan.png',
       'Gambar 6. Perbandingan akurasi enam algoritma deteksi ritual terhadap tingkat galat GPS (σ). '
       "Sa'i, Geofence Miqat, dan Deteksi Arafah bertahan di atas 99% pada seluruh σ, sedangkan penghitung "
       'Tawaf dan deteksi Jamarat menurun pada σ ≥ 10 m.')
heading('4.8', 'Keterbatasan', level=2)
para('Pertama, galat sistematik Haversine: asumsi bumi bola menimbulkan galat sekitar 0,2–0,5% terhadap '
     'elipsoid; dapat diperkecil dengan Vincenty bila presisi jarak jauh diperlukan. Kedua, sensitivitas '
     'Tawaf pada gangguan besar: radius tawaf yang kecil (25 m) dan deteksi persilangan sudut membuat '
     'penghitung rawan over-count pada σ ≥ 15 m; ambang debounce dan penghalusan lintasan (misalnya tapis '
     'Kalman) berpotensi memperbaikinya. Ketiga, kasus tepi garis Hajar Aswad: deteksi persilangan gagal '
     'bila sebuah sampel jatuh tepat pada sudut 0°; dalam praktik peristiwa ini berpeluang nol karena '
     'gangguan GPS, namun merupakan celah logika yang perlu ditangani secara eksplisit. Keempat, '
     'non-deteksi Jamarat pada σ tinggi akibat radius 30 m; perlu penyesuaian radius adaptif. Kelima, '
     'penyederhanaan batas: poligon Arafah 5 titik dan Masjid Namirah sebagai titik tunggal dengan radius '
     '200 m adalah aproksimasi dari batas sebenarnya. Keenam, kesenjangan simulasi–lapangan: model Gaussian '
     'isotropik tidak menangkap multipath dan bias sistematik akibat menara tinggi di sekitar Masjidil '
     'Haram; validasi lapangan dengan jejak GPS nyata direkomendasikan sebagai penelitian lanjutan.')

# ================= 5. SIMPULAN =================
heading('5', 'SIMPULAN')
para('Penelitian ini menjawab satu pertanyaan yang selama ini dilewatkan oleh pengembangan aplikasi '
     'pendamping haji: seberapa jauh deteksi ritual berbasis geometri GPS dapat dipercaya ketika posisi '
     'yang diterimanya salah. Enam algoritma pada sistem Mabrur diuji melalui simulasi Monte Carlo dengan '
     'injeksi derau Gaussian pada enam tingkat galat (σ = 0 sampai 15 m), menggunakan kode produksi '
     'aplikasi tanpa modifikasi logika. Seluruh algoritma terbukti benar pada kondisi ideal (σ = 0, akurasi '
     '100%), sehingga setiap penurunan kinerja yang teramati murni disebabkan galat posisi, bukan cacat '
     'logika.')
para('Hasil pengujian memperlihatkan tiga tingkatan ketahanan. Pada tingkatan pertama, penghitung sa’i '
     'bertahan sempurna (100%) hingga galat terburuk, dan geofence miqat maupun deteksi batas Arafah tetap '
     'di atas 99,4% — kesalahan hanya muncul pada pita sempit di sekitar garis batas. Pada tingkatan kedua, '
     'penghitung tawaf andal penuh hingga σ = 10 m namun turun ke 72,67% pada σ = 15 m dengan kecenderungan '
     'kelebihan hitung di sekitar garis Hajar Aswad. Pada tingkatan ketiga, identifikasi pilar jamarat '
     'akurat penuh hingga σ = 5 m lalu melemah menjadi 83,53% pada σ = 15 m; yang perlu digarisbawahi, '
     'kesalahan identifikasi antar-pilar nyaris tidak terjadi (maksimum 0,32%) — penurunan hampir '
     'seluruhnya berupa kegagalan deteksi karena posisi terukur terdorong keluar radius 30 m. Pola ketiga '
     'tingkatan ini mengerucut pada satu temuan umum: ketahanan algoritma geometris ditentukan oleh rasio '
     'antara separasi geometris objek ritual dan besar galat posisi. Jarak Safa–Marwah yang 419 m membuat '
     'sa’i praktis kebal terhadap galat GPS ponsel; radius tawaf yang hanya 25 m menjadikannya komponen '
     'paling rawan.')
para('Temuan tersebut membawa implikasi praktis yang langsung dapat dipakai perancang aplikasi sejenis. '
     'Karena galat GPS ponsel di lingkungan padat bangunan berada pada kisaran 7–13 m, seluruh algoritma '
     'yang diuji berada dalam wilayah kerja yang aman untuk kondisi lapangan tipikal — kecuali pada '
     'skenario terburuk di sekitar Masjidil Haram, tempat penghitung tawaf dan deteksi jamarat memerlukan '
     'mitigasi tambahan. Mitigasi yang paling menjanjikan mengikuti arah temuan: penghalusan lintasan '
     '(misalnya tapis Kalman) sebelum deteksi persilangan sudut untuk tawaf, dan radius deteksi adaptif '
     'atau pemanfaatan konteks urutan lempar harian untuk jamarat.')
para('Penelitian ini dibatasi oleh sifat simulasinya: model derau Gaussian isotropik tidak menangkap bias '
     'multipath yang sesungguhnya terjadi di antara menara-menara sekitar Masjidil Haram, dan lintasan '
     'simulasi tidak memuat perilaku jamaah yang berhenti, berbalik, atau terbawa arus kerumunan. Validasi '
     'lapangan dengan rekaman jejak GPS nyata dari musim haji menjadi kelanjutan yang paling mendesak. Dua '
     'arah lanjutan lain telah disiapkan pada sistem yang sama: evaluasi fitur keselamatan jamaah '
     '(peringatan SOS dan pemantauan kelompok) serta evaluasi arsitektur luring-pertama untuk kondisi '
     'jaringan yang padat — keduanya melengkapi validasi algoritmik ini menuju penilaian sistem secara '
     'utuh.')

# ================= DAFTAR PUSTAKA =================
heading(None, 'DAFTAR PUSTAKA')
REFS = [
    'Katadata Databoks, “Kuota Haji 2025 untuk Jemaah Reguler, Lansia, Pembimbing, dan Petugas Daerah,” '
    '2025. [Daring]. [Lengkapi URL & tanggal akses; utamakan dokumen resmi Kemenag RI].',
    'NU Online, “Ada 45 Ribu Jamaah Haji Lansia Tahun 2024, Ini Upaya Pemerintah,” 2024. [Daring]. '
    '[Lengkapi URL & tanggal akses].',
    '[Penulis — lengkapi], “Aplikasi Pembelajaran Manasik Haji dan Umroh Berbasis Android,” Bitnet: '
    'Jurnal Pendidikan Teknologi Informasi, Universitas Muhammadiyah Palangkaraya. [Lengkapi vol/no/tahun].',
    '[Penulis — lengkapi], “Buku Saku Ibadah Manasik Haji dan Umroh Berbasis Android.” [Lengkapi jurnal, '
    'vol/no/tahun].',
    '[Penulis — lengkapi], “Pemodelan dan Implementasi Aplikasi Mobile Umrah Guide Menggunakan Unified '
    'Modeling Language,” Jurnal Sains dan Informatika, Politeknik Negeri Tanah Laut. [Lengkapi vol/no/tahun].',
    'T. Budiawan dkk., “Development of Android Based Hajj and Umrah Pilgrims Monitoring Application in '
    'Dago Wisata International,” IJCCS (Indonesian Journal of Computing and Cybernetics Systems), '
    'Universitas Gadjah Mada. [Lengkapi vol/no/halaman/tahun/DOI].',
    '[Penulis — lengkapi], “Automatic Hajj and Umrah Ritual Detection Using IMU Sensors,” King Fahd '
    'University of Petroleum & Minerals, ±2022. [Lengkapi jurnal/prosiding & DOI].',
    '[Penulis — lengkapi], “Autonomous Mobile-Based Model for Tawaf/Sa’ay Rounds Counting with Supported '
    'Supplications from the Quran and Sunna’a,” 2023. [Lengkapi jurnal & DOI].',
    '[Penulis — lengkapi], “Hajj Crowd Management Using CNN-Based Approach,” 2021. [Lengkapi jurnal & DOI].',
    'P. A. Zandbergen dan S. J. Barbeau, “Positional Accuracy of Assisted GPS Data from High-Sensitivity '
    'GPS-Enabled Mobile Phones,” Journal of Navigation, vol. 64, no. 3, hlm. 381–399, 2011. [Verifikasi halaman].',
    'K. Merry dan P. Bettinger, “Smartphone GPS accuracy study in an urban environment,” PLOS ONE, '
    'vol. 14, no. 7, e0219890, 2019.',
    '[Penulis — lengkapi], “Accuracy and Reliability of Smartphones GNSS Applications: A Comparative Study '
    'with Handheld GPS Device for Position Determination,” World Scientific News, vol. 202, 2025. '
    '[Verifikasi kelayakan sumber; ganti bila kurang kuat].',
    'R. W. Sinnott, “Virtues of the Haversine,” Sky & Telescope, vol. 68, no. 2, hlm. 159, 1984.',
    'T. Vincenty, “Direct and Inverse Solutions of Geodesics on the Ellipsoid with Application of Nested '
    'Equations,” Survey Review, vol. 23, no. 176, hlm. 88–93, 1975.',
    'M. Shimrat, “Algorithm 112: Position of point relative to polygon,” Communications of the ACM, '
    'vol. 5, no. 8, hlm. 434, 1962.',
]
for i, ref in enumerate(REFS, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)  # hanging indent
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f'[{i}] {ref}')
    r.font.size = Pt(10)

doc.save(OUT)

# verifikasi ringkas
check = Document(OUT)
print(f'[OK] {OUT}')
print(f'     paragraf={len(check.paragraphs)}, tabel={len(check.tables)}, gambar={len(check.inline_shapes)}')
print(f'     ukuran={os.path.getsize(OUT) / 1024:.0f} KB')
